from pathlib import Path

from openpyxl import load_workbook
import docx


# from https://github.com/python-openxml/python-docx/issues/25#issuecomment-400787031
def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def parse_sheets(xlsx):
    wb = load_workbook(xlsx)
    columns = [col for col in wb.active.iter_cols()]
    header, data = [c.value for c in columns[1]], [[c.value for c in col] for col in columns[2:]]
    organized = []
    for col in data:
        question = col[0]
        answers = {}
        for i in range(len(col)):
            if i == 0:
                continue
            answers[header[i]] = col[i]
        organized.append((question, answers))

    return organized


def export_docx(data, out_file):
    doc = docx.Document()
    for question, answers in data:
        doc.add_heading(question, level=1)
        prev = None
        for answerer, answer in answers.items():
            if not answer and not answerer:
                continue
            par = doc.add_paragraph('', style='List Number')
            par.add_run(answerer, 'Emphasis').font.bold = True
            par.add_run(' — ')
            par.add_run(str(answer))
            if prev:
                list_number(doc, par, prev=prev, num=True)
            else:
                list_number(doc, par, num=True)
            prev = par

    doc.save(out_file)


if __name__ == '__main__':
    in_path = 'input'
    out_path = 'output'
    for f in Path(in_path).glob('*.xlsx'):
        print(f'parsing {f}')
        parsed = parse_sheets(f)
        out_file = Path(out_path) / (f.stem + '.docx')
        export_docx(parsed, out_file)
